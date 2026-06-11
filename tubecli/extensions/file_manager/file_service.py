"""
File Service — Sandboxed file operations.
Security: Only allows operations within allowed directories.
"""
import os
import shutil
import glob
import time
from pathlib import Path
from typing import Dict, List, Optional, Any


# Allowed root directories (user can customize)
DEFAULT_ALLOWED_ROOTS = [
    os.path.expanduser("~/Desktop"),
    os.path.expanduser("~/Documents"),
    os.path.expanduser("~/Downloads"),
]

# Dangerous paths that are NEVER allowed
BLOCKED_PATHS = [
    "C:\\Windows", "C:\\Program Files", "C:\\Program Files (x86)",
    "/usr", "/bin", "/sbin", "/etc", "/boot", "/lib",
    os.path.expanduser("~/.ssh"),
    os.path.expanduser("~/AppData/Local"),
]

MAX_FILE_SIZE_MB = 50  # Max file size for read operations


class FileService:
    """Sandboxed file operations service."""

    def __init__(self, extra_roots: Optional[List[str]] = None):
        self.allowed_roots = list(DEFAULT_ALLOWED_ROOTS)
        # Add data dir
        data_dir = os.environ.get("TUBECLI_DATA_DIR", "data")
        self.allowed_roots.append(os.path.abspath(data_dir))
        if extra_roots:
            self.allowed_roots.extend(extra_roots)
        # Normalize all paths
        self.allowed_roots = [os.path.normpath(r) for r in self.allowed_roots]

    def _validate_path(self, path: str) -> str:
        """Validate and normalize path. Raises ValueError if blocked."""
        # Expand ~ and environment variables
        expanded = os.path.expanduser(os.path.expandvars(path))
        normalized = os.path.normpath(os.path.abspath(expanded))

        # Check blocked paths
        for blocked in BLOCKED_PATHS:
            blocked_norm = os.path.normpath(blocked)
            if normalized.startswith(blocked_norm):
                raise ValueError(f"Đường dẫn bị chặn vì lý do bảo mật: {path}")

        # Check if within allowed roots
        in_allowed = any(
            normalized.startswith(os.path.normpath(root))
            for root in self.allowed_roots
        )
        if not in_allowed:
            raise ValueError(
                f"Đường dẫn nằm ngoài vùng cho phép: {path}\n"
                f"Vùng cho phép: {', '.join(self.allowed_roots)}"
            )

        return normalized

    def _file_info(self, path: str) -> Dict[str, Any]:
        """Get file/folder info dict."""
        stat = os.stat(path)
        return {
            "name": os.path.basename(path),
            "path": path,
            "is_dir": os.path.isdir(path),
            "size": stat.st_size if os.path.isfile(path) else 0,
            "size_human": self._human_size(stat.st_size) if os.path.isfile(path) else "",
            "modified": time.strftime("%Y-%m-%d %H:%M:%S", time.localtime(stat.st_mtime)),
            "created": time.strftime("%Y-%m-%d %H:%M:%S", time.localtime(stat.st_ctime)),
            "extension": os.path.splitext(path)[1].lower() if os.path.isfile(path) else "",
        }

    @staticmethod
    def _human_size(size_bytes: int) -> str:
        """Convert bytes to human readable string."""
        for unit in ("B", "KB", "MB", "GB"):
            if size_bytes < 1024:
                return f"{size_bytes:.1f} {unit}"
            size_bytes /= 1024
        return f"{size_bytes:.1f} TB"

    # ── Core Operations ──────────────────────────────────────

    def list_dir(self, path: str, show_hidden: bool = False) -> Dict[str, Any]:
        """List contents of a directory."""
        safe_path = self._validate_path(path)
        if not os.path.isdir(safe_path):
            raise FileNotFoundError(f"Thư mục không tồn tại: {path}")

        items = []
        for entry in sorted(os.listdir(safe_path)):
            if not show_hidden and entry.startswith("."):
                continue
            full = os.path.join(safe_path, entry)
            try:
                items.append(self._file_info(full))
            except (PermissionError, OSError):
                items.append({"name": entry, "path": full, "error": "Permission denied"})

        return {
            "path": safe_path,
            "parent": os.path.dirname(safe_path),
            "items": items,
            "count": len(items),
            "dirs": sum(1 for i in items if i.get("is_dir")),
            "files": sum(1 for i in items if not i.get("is_dir") and "error" not in i),
        }

    def create_folder(self, path: str) -> Dict[str, Any]:
        """Create a folder (including parents)."""
        safe_path = self._validate_path(path)
        if os.path.exists(safe_path):
            return {"status": "exists", "path": safe_path, "message": f"Thư mục đã tồn tại: {path}"}
        os.makedirs(safe_path, exist_ok=True)
        return {"status": "created", "path": safe_path, "message": f"Đã tạo thư mục: {safe_path}"}

    def create_file(self, path: str, content: str = "") -> Dict[str, Any]:
        """Create a file with optional content. Supports txt, docx, xlsx."""
        safe_path = self._validate_path(path)
        parent_dir = os.path.dirname(safe_path)
        os.makedirs(parent_dir, exist_ok=True)

        ext = os.path.splitext(safe_path)[1].lower()
        
        try:
            if ext == '.docx':
                from docx import Document
                doc = Document()
                if content:
                    doc.add_paragraph(content)
                doc.save(safe_path)
                size = os.path.getsize(safe_path)
            elif ext == '.xlsx':
                from openpyxl import Workbook
                wb = Workbook()
                ws = wb.active
                if content:
                    for row_idx, line in enumerate(content.split('\n'), 1):
                        cells = line.split('\t') if '\t' in line else line.split(',')
                        for col_idx, cell in enumerate(cells, 1):
                            ws.cell(row=row_idx, column=col_idx, value=cell.strip())
                wb.save(safe_path)
                size = os.path.getsize(safe_path)
            else:
                mode = "w" if not os.path.exists(safe_path) else "w"
                with open(safe_path, mode, encoding="utf-8") as f:
                    f.write(content)
                size = os.path.getsize(safe_path)
                
            return {
                "status": "created",
                "path": safe_path,
                "size": size,
                "extension": ext,
                "message": f"Đã tạo file: {safe_path}",
            }
        except ImportError as e:
            raise RuntimeError(f"Thiếu thư viện để xử lý file {ext}: {str(e)}")
        except Exception as e:
            raise RuntimeError(f"Lỗi tạo file {ext}: {str(e)}")

    def read_file(self, path: str, max_lines: int = 1000) -> Dict[str, Any]:
        """Read file content (supports txt, docx, xlsx) and metadata."""
        safe_path = self._validate_path(path)
        if not os.path.isfile(safe_path):
            raise FileNotFoundError(f"File không tồn tại: {path}")

        size = os.path.getsize(safe_path)
        if size > MAX_FILE_SIZE_MB * 1024 * 1024:
            return {"error": f"File quá lớn ({self._human_size(size)}). Giới hạn: {MAX_FILE_SIZE_MB}MB"}

        ext = os.path.splitext(safe_path)[1].lower()
        content = ""
        lines = []

        try:
            if ext == '.docx':
                from docx import Document
                doc = Document(safe_path)
                for i, p in enumerate(doc.paragraphs):
                    if i >= max_lines: break
                    if p.text.strip():
                        lines.append(p.text)
                content = "\n".join(lines)
            elif ext == '.xlsx':
                from openpyxl import load_workbook
                wb = None
                try:
                    wb = load_workbook(safe_path, read_only=True, data_only=True)
                    for ws in wb.worksheets:
                        lines.append(f"--- Sheet: {ws.title} ---")
                        for i, row in enumerate(ws.iter_rows(values_only=True)):
                            if len(lines) >= max_lines: break
                            row_data = [str(c) if c is not None else "" for c in row]
                            if any(row_data):
                                lines.append("\t".join(row_data))
                        if len(lines) >= max_lines: break
                    content = "\n".join(lines)
                finally:
                    if wb:
                        wb.close()
            else:
                with open(safe_path, "r", encoding="utf-8") as f:
                    for i, line in enumerate(f):
                        if i >= max_lines:
                            break
                        lines.append(line.rstrip("\n"))
                content = "\n".join(lines)
        except UnicodeDecodeError:
            return {
                "path": safe_path,
                "is_binary": True,
                "size": size,
                "size_human": self._human_size(size),
                "extension": ext,
                "modified": time.strftime("%Y-%m-%d %H:%M:%S", time.localtime(os.path.getmtime(safe_path))),
                "message": "File nhị phân, không thể đọc dạng text.",
            }
        except ImportError as e:
            return {"error": f"Thiếu thư viện để đọc file {ext}: {str(e)}"}
        except Exception as e:
            return {"error": f"Lỗi đọc file: {str(e)}"}

        return {
            "path": safe_path,
            "content": content,
            "lines": len(lines),
            "truncated": len(lines) >= max_lines,
            "size": size,
            "size_human": self._human_size(size),
            "extension": ext,
            "modified": time.strftime("%Y-%m-%d %H:%M:%S", time.localtime(os.path.getmtime(safe_path))),
            "created": time.strftime("%Y-%m-%d %H:%M:%S", time.localtime(os.path.getctime(safe_path))),
        }

    def delete(self, path: str) -> Dict[str, Any]:
        """Delete a file or folder."""
        safe_path = self._validate_path(path)
        if not os.path.exists(safe_path):
            raise FileNotFoundError(f"Đường dẫn không tồn tại: {path}")

        if os.path.isdir(safe_path):
            item_count = sum(len(files) + len(dirs) for _, dirs, files in os.walk(safe_path))
            shutil.rmtree(safe_path)
            return {"status": "deleted", "path": safe_path, "type": "folder", "items_removed": item_count}
        else:
            os.remove(safe_path)
            return {"status": "deleted", "path": safe_path, "type": "file"}

    def move(self, src: str, dst: str) -> Dict[str, Any]:
        """Move or rename a file/folder."""
        safe_src = self._validate_path(src)
        safe_dst = self._validate_path(dst)
        if not os.path.exists(safe_src):
            raise FileNotFoundError(f"Nguồn không tồn tại: {src}")

        # If dst is a directory, move into it
        if os.path.isdir(safe_dst):
            safe_dst = os.path.join(safe_dst, os.path.basename(safe_src))

        os.makedirs(os.path.dirname(safe_dst), exist_ok=True)
        shutil.move(safe_src, safe_dst)
        return {"status": "moved", "from": safe_src, "to": safe_dst}

    def copy(self, src: str, dst: str) -> Dict[str, Any]:
        """Copy a file or folder."""
        safe_src = self._validate_path(src)
        safe_dst = self._validate_path(dst)
        if not os.path.exists(safe_src):
            raise FileNotFoundError(f"Nguồn không tồn tại: {src}")

        os.makedirs(os.path.dirname(safe_dst), exist_ok=True)

        if os.path.isdir(safe_src):
            shutil.copytree(safe_src, safe_dst)
        else:
            shutil.copy2(safe_src, safe_dst)
        return {"status": "copied", "from": safe_src, "to": safe_dst}

    def info(self, path: str) -> Dict[str, Any]:
        """Get detailed file/folder info."""
        safe_path = self._validate_path(path)
        if not os.path.exists(safe_path):
            raise FileNotFoundError(f"Đường dẫn không tồn tại: {path}")

        info = self._file_info(safe_path)
        if info["is_dir"]:
            # Count contents
            total_files = 0
            total_size = 0
            for root, dirs, files in os.walk(safe_path):
                total_files += len(files)
                for f in files:
                    try:
                        total_size += os.path.getsize(os.path.join(root, f))
                    except OSError:
                        pass
            info["total_files"] = total_files
            info["total_size"] = total_size
            info["total_size_human"] = self._human_size(total_size)
        return info

    def search(self, path: str, pattern: str = "*", recursive: bool = True) -> Dict[str, Any]:
        """Search files matching pattern (supports mutiple patterns separated by |)."""
        safe_path = self._validate_path(path)
        if not os.path.isdir(safe_path):
            raise FileNotFoundError(f"Thư mục không tồn tại: {path}")

        matches = []
        patterns = pattern.split("|") if "|" in pattern else [pattern]
        
        seen_paths = set()
        for p in patterns:
            search_pattern = os.path.join(safe_path, "**", p.strip()) if recursive else os.path.join(safe_path, p.strip())
            for match in glob.iglob(search_pattern, recursive=recursive):
                if match in seen_paths: continue
                seen_paths.add(match)
                try:
                    matches.append(self._file_info(match))
                except (PermissionError, OSError):
                    continue
                if len(matches) >= 200:
                    break
            if len(matches) >= 200:
                break

        return {"path": safe_path, "pattern": pattern, "matches": matches, "count": len(matches)}

    def get_allowed_roots(self) -> List[Dict[str, str]]:
        """Return list of allowed root directories for the UI."""
        roots = []
        for r in self.allowed_roots:
            if os.path.isdir(r):
                roots.append({"path": r, "name": os.path.basename(r) or r, "exists": True})
            else:
                roots.append({"path": r, "name": os.path.basename(r) or r, "exists": False})
        return roots


# Global singleton
file_service = FileService()
